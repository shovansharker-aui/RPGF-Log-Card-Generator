from pathlib import Path
import pandas as pd

from docx import Document
from docxcompose.composer import Composer

# ==========================================================
# PROJECT PATHS
# ==========================================================

BASE_DIR = Path(__file__).parent

TEMPLATE = BASE_DIR / "Template.docx"
EXCEL = BASE_DIR / "Equipment.xlsx"
OUTPUT = BASE_DIR / "output"

OUTPUT.mkdir(exist_ok=True)

# ==========================================================
# REPLACE TEXT WHILE PRESERVING FORMATTING
# ==========================================================

def replace_in_paragraph(paragraph, old_text, new_text):
    """
    Replace text while preserving formatting.
    Works when the placeholder exists within a single run.
    """

    if old_text not in paragraph.text:
        return

    for run in paragraph.runs:
        if old_text in run.text:
            run.text = run.text.replace(old_text, str(new_text))


def replace_everywhere(doc, old_text, new_text):

    # Paragraphs
    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph, old_text, new_text)

    # Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_paragraph(paragraph, old_text, new_text)


# ==========================================================
# CREATE ONE DOCUMENT
# ==========================================================

def create_document(row, month1, month2, month3):

    doc = Document(TEMPLATE)

    # Equipment Information
    replace_everywhere(doc, "(Enter Equipment Name)", row["Equipment Name"])
    replace_everywhere(doc, "(Enter Equipment ID)", row["Equipment No."])
    replace_everywhere(doc, "(Enter Equipment Location)", row["Location"])

    # Month Headers
    replace_everywhere(doc, "(Enter Month1’Year)", month1)
    replace_everywhere(doc, "(Enter Month2’Year)", month2)
    replace_everywhere(doc, "(Enter Month3’Year)", month3)

    return doc


# ==========================================================
# MAIN
# ==========================================================

def main():

    print("=" * 70)
    print("RENATA PLC")
    print("Preventive Maintenance Log Card Generator")
    print("=" * 70)

    # ------------------------------------------------------
    # Read Equipment Sheet
    # ------------------------------------------------------

    equipment_df = pd.read_excel(
        EXCEL,
        sheet_name="Equipment"
    )

    # ------------------------------------------------------
    # Read Settings Sheet
    # ------------------------------------------------------

    settings_df = pd.read_excel(
        EXCEL,
        sheet_name="Settings"
    )

    month1 = settings_df.loc[0, "Month1"]
    month2 = settings_df.loc[0, "Month2"]
    month3 = settings_df.loc[0, "Month3"]

    print()
    print("Schedule")
    print("-------------------------------")
    print(f"Month 1 : {month1}")
    print(f"Month 2 : {month2}")
    print(f"Month 3 : {month3}")
    print()

    print(f"Total Equipment : {len(equipment_df)}")
    print()

    # ------------------------------------------------------
    # First Document
    # ------------------------------------------------------

    master = create_document(
        equipment_df.iloc[0],
        month1,
        month2,
        month3
    )

    composer = Composer(master)

    print(f"Added : {equipment_df.iloc[0]['Equipment Name']}")

    # ------------------------------------------------------
    # Remaining Documents
    # ------------------------------------------------------

    for index in range(1, len(equipment_df)):

        row = equipment_df.iloc[index]

        doc = create_document(
            row,
            month1,
            month2,
            month3
        )

        composer.append(doc)

        print(f"Added : {row['Equipment Name']}")

    # ------------------------------------------------------
    # Save Final Document
    # ------------------------------------------------------

    output_file = OUTPUT / "Preventive_Maintenance_Log_Cards.docx"

    composer.save(output_file)

    print()
    print("=" * 70)
    print("SUCCESS")
    print("=" * 70)
    print(f"Combined document created successfully.")
    print(output_file)
    print("=" * 70)


# ==========================================================
# START
# ==========================================================

if __name__ == "__main__":
    main()