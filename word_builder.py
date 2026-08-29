"""
word_builder.py

Word document helper
"""

from docx import Document


class WordBuilder:

    # ==================================================
    # Constructor
    # ==================================================

    def __init__(self, template_path):

        self.template_path = template_path

        self.doc = None

    # ==================================================
    # Load Template
    # ==================================================

    def load(self):

        self.doc = Document(self.template_path)

    # ==================================================
    # Save Document
    # ==================================================

    def save(self, filename):

        self.doc.save(filename)

    # ==================================================
    # Replace Placeholder in Paragraph
    # ==================================================

    def replace_paragraph(self, paragraph, old, new):

        if old not in paragraph.text:

            return

        new_text = paragraph.text.replace(

            old,

            str(new)

        )

        if paragraph.runs:

            paragraph.runs[0].text = new_text

            for run in paragraph.runs[1:]:

                run.text = ""

        else:

            paragraph.text = new_text

    # ==================================================
    # Replace in Tables
    # ==================================================

    def replace_tables(self, old, new):

        for table in self.doc.tables:

            for row in table.rows:

                for cell in row.cells:

                    for paragraph in cell.paragraphs:

                        self.replace_paragraph(

                            paragraph,

                            old,

                            new

                        )

    # ==================================================
    # Replace in Document
    # ==================================================

    def replace(self, old, new):

        for paragraph in self.doc.paragraphs:

            self.replace_paragraph(

                paragraph,

                old,

                new

            )

        self.replace_tables(

            old,

            new

        )

    # ==================================================
    # Replace Multiple Placeholders
    # ==================================================

    def replace_all(self, placeholders):

        for placeholder, value in placeholders.items():

            self.replace(

                placeholder,

                value

            )

    # ==================================================
    # Document Information
    # ==================================================

    def paragraph_count(self):

        return len(self.doc.paragraphs)

    def table_count(self):

        return len(self.doc.tables)